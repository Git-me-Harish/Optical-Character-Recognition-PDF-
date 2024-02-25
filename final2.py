import streamlit as st
import fitz 
from unidecode import unidecode
import pandas as pd
import numpy as np
import re
from docx import Document 
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import tempfile
import os
import base64

def extract_headings_and_content(pdf_path):
    block_dict = {}
    headings_list = []
    text_list = []
    
    doc = fitz.open(pdf_path)
    
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        file_dict = page.get_text('dict') 
        block = file_dict['blocks'] 
        block_dict[page_num] = block 

    spans = pd.DataFrame(columns=['xmin', 'ymin', 'xmax', 'ymax', 'text', 'tag'])
    rows = []
    for page_num, blocks in block_dict.items():
        for block in blocks:
            if block['type'] == 0:
                for line in block['lines']:
                    for span in line['spans']:
                        xmin, ymin, xmax, ymax = list(span['bbox'])
                        font_size = span['size']
                        text = unidecode(span['text'])
                        span_font = span['font']
                        is_upper = False
                        is_bold = False
                        if "bold" in span_font.lower():
                            is_bold = True
                        if re.sub("[\(\[].*?[\)\]]", "", text).isupper():
                            is_upper = True
                        if text.replace(" ","") !=  "":
                            rows.append((xmin, ymin, xmax, ymax, text, is_upper, is_bold, span_font, font_size))
    span_df = pd.DataFrame(rows, columns=['xmin','ymin','xmax','ymax', 'text', 'is_upper','is_bold','span_font', 'font_size'])

    span_scores = []
    for index, span_row in span_df.iterrows():
        score = round(span_row.font_size)
        text = span_row.text
        if not re.search('[(_:/,#%\=@)]', text):
            if span_row.is_bold:
                score += 1 
            if span_row.is_upper:
                score += 1
        span_scores.append(score)

    values, counts = np.unique(span_scores, return_counts=True)
    style_dict = {}
    for value, count in zip(values, counts):
        style_dict[value] = count

    p_size = max(style_dict, key=style_dict.get)
    idx = 0
    tag = {}
    for size in sorted(values, reverse=True):
        idx += 1
        if size == p_size:
            idx = 0
            tag[size] = 'p'
        if size > p_size:
            idx += 1
            tag[size] = f'h{idx}'
        if size < p_size:
            idx += 1
            tag[size] = f's{idx}'

    span_tags = [tag[score] for score in span_scores]
    span_df['tag'] = span_tags

    tmp = []
    heading = ''                                                                                                                
    for index, span_row in span_df.iterrows():
        text = span_row.text
        tag = span_row.tag
        if 'h' in tag:
            if tmp:
                text_list.append('\n'.join(tmp))
                tmp = []
            if heading and text.startswith(heading.split()[-1]):
                # Combine the current heading with the previous one
                headings_list[-1] = ' '.join([headings_list[-1], text])
            else:
                headings_list.append(text)
            heading = text
        else:
            tmp.append(text)
    text_list.append('\n'.join(tmp))
    text_list = text_list[1:]
    
    text_df = pd.DataFrame(zip(headings_list, text_list), columns=['heading', 'content'])
    
    return text_df

def save_to_word(text_df):
    # Create a temporary directory to store the Word document
    with tempfile.TemporaryDirectory() as temp_dir:
        output_path = os.path.join(temp_dir, 'extracted_content.docx')
        
        document = Document()
        
        for index, row in text_df.iterrows():
            heading = row['heading']
            content = row['content']
            
            # Add heading to the document with proper alignment
            document.add_heading(heading, level=1)
            
            # Add content to the document with proper alignment
            p = document.add_paragraph(content)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            
            # Add space between headings and content
            document.add_paragraph('\n')
        
        # Save the document to the specified output path
        document.save(output_path)
        
        # Read the saved file as bytes
        with open(output_path, "rb") as f:
            file_bytes = f.read()
        
        return file_bytes

def main():
    st.title("PDF Heading and Content Extraction")
    st.sidebar.title("Upload PDF")

    # Allow user to upload PDF file
    uploaded_file = st.sidebar.file_uploader("Choose a PDF file", type="pdf")

    if uploaded_file is not None:
        # Save the uploaded file to a temporary directory
        temp_dir = tempfile.TemporaryDirectory()
        pdf_path = os.path.join(temp_dir.name, uploaded_file.name)
        with open(pdf_path, "wb") as f:
            f.write(uploaded_file.getvalue())
        
        st.success("PDF file uploaded successfully!")
        headings_content_df = extract_headings_and_content(pdf_path)

        # Display headings in a dropdown
        selected_heading = st.selectbox("Select a heading:", headings_content_df['heading'].tolist())

        # Display content of selected heading
        st.subheader(f"Content under '{selected_heading}':")
        content = headings_content_df[headings_content_df['heading'] == selected_heading]['content'].values[0]
        st.write(content)

        # Button to save to Word
        if st.button("Save to Word"):
            file_bytes = save_to_word(headings_content_df)
            b64 = base64.b64encode(file_bytes).decode()
            href = f'<a href="data:file/docx;base64,{b64}" download="extracted_content.docx">Download Extracted Texts</a>'
            st.markdown(href, unsafe_allow_html=True)

if __name__=="__main__":
    main()
