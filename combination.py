import streamlit as st
import fitz  # PyMuPDF for parsing PDF
import os
from docx import Document

# Function to extract headings and their content from PDF document
def extract_headings_and_content(pdf_file):
    headings_content = {}
    with fitz.open(pdf_file) as doc:
        for page in doc:
            output = page.get_text("blocks")
            previous_block_id = 0
            heading = ""
            content = ""
            for block in output:
                if block[6] == 0:  # We only take the text
                    if previous_block_id != block[5]:  # Compare the block number
                        if heading:  # If heading is not empty, store previous heading and content
                            headings_content[heading] = content
                        heading = block[4].strip()
                        content = block[4].strip()  # Initialize content with heading text
                    else:
                        content += "\n" + block[4].strip()  # Append text to content
                    previous_block_id = block[5]
            # Store last heading and content pair
            if heading:
                headings_content[heading] = content
    return headings_content

# Function to download text content as a Word file
def download_word(content, filename):
    doc = Document()
    doc.add_heading('Extracted Content', 0)
    for heading, text in content.items():
        doc.add_heading(heading, level=1)
        doc.add_paragraph(text)
    doc.save(filename)

# Main function for PDF Heading Extraction
def pdf_heading_extraction():
    st.title("PDF Heading Extraction")

    # Allow users to upload PDF document
    uploaded_file = st.file_uploader("Upload PDF document", type=["pdf"])

    if uploaded_file:
        st.success("PDF uploaded successfully!")

        # Extract headings and their content from PDF
        temp_file_path = "temp_uploaded_file.pdf"
        with open(temp_file_path, "wb") as temp_file:
            temp_file.write(uploaded_file.getbuffer())

        headings_content = extract_headings_and_content(temp_file_path)

        # Display extracted headings in dropdown menu
        if len(headings_content) > 0:
            selected_heading = st.selectbox("Select a heading:", list(headings_content.keys()))

            # Display content of selected heading
            st.subheader(f"Content under '{selected_heading}':")
            st.write(headings_content[selected_heading])

            # Download buttons
            st.markdown("<hr>", unsafe_allow_html=True)
            st.subheader("Download")
            st.write("Click the button below to download:")
            st.download_button(f"Download All Content (PDF)", temp_file_path)
            st.download_button(f"Download '{selected_heading}' Content (Text)", headings_content[selected_heading], f"{selected_heading}.txt")
            st.download_button(f"Download All Content (Word)", headings_content, "extracted_content.docx")

        else:
            st.warning("No headings found in the uploaded PDF.")

        # Remove temporary file
        os.remove(temp_file_path)

# Main function to create Streamlit interface
def main():
    st.set_page_config(page_title="PDF Heading Extraction", page_icon=":notebook:", layout="wide")

    st.sidebar.title("Menu")
    menu_option = st.sidebar.radio(
        "Select Option",
        ["PDF Heading Extraction"]
    )

    if menu_option == "PDF Heading Extraction":
        pdf_heading_extraction()

if __name__ == "__main__":
    main()
